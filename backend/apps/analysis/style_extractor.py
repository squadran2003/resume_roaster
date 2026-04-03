"""
Extract visual style information from the original resume (PDF or DOCX)
so the rewritten PDF download can closely match the original formatting.
"""

import io
import logging
import re
from collections import Counter

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Font-family mapping: PDF/DOCX font names -> CSS font-family stacks
# ---------------------------------------------------------------------------
_FONT_MAP = {
    "times": '"Times New Roman", Times, serif',
    "arial": "Arial, Helvetica, sans-serif",
    "helvetica": "Arial, Helvetica, sans-serif",
    "calibri": "Calibri, sans-serif",
    "garamond": "Garamond, serif",
    "georgia": "Georgia, serif",
    "cambria": "Cambria, serif",
    "palatino": '"Palatino Linotype", Palatino, serif',
    "century": '"Century Gothic", sans-serif',
    "verdana": "Verdana, sans-serif",
    "tahoma": "Tahoma, sans-serif",
    "trebuchet": '"Trebuchet MS", sans-serif',
    "lucida": '"Lucida Sans", sans-serif',
    "courier": '"Courier New", Courier, monospace',
    "roboto": "Roboto, sans-serif",
    "lato": "Lato, sans-serif",
    "opensans": '"Open Sans", sans-serif',
    "open sans": '"Open Sans", sans-serif',
    "montserrat": "Montserrat, sans-serif",
    "raleway": "Raleway, sans-serif",
    "merriweather": "Merriweather, serif",
    "poppins": "Poppins, sans-serif",
    "nunito": "Nunito, sans-serif",
    "source": '"Source Sans Pro", sans-serif',
    "inter": "Inter, sans-serif",
    "avenir": "Avenir, sans-serif",
    "futura": "Futura, sans-serif",
    "gill": '"Gill Sans", sans-serif',
    "optima": "Optima, sans-serif",
    "book antiqua": '"Book Antiqua", Palatino, serif',
    "rockwell": "Rockwell, serif",
}


def _map_font(raw_name: str) -> str:
    """Map a PDF/DOCX font name to a CSS font-family value."""
    # Strip subset prefix (e.g. "ABCDEF+")
    clean = re.sub(r"^[A-Z]{6}\+", "", raw_name)
    lower = clean.lower().replace("-", "").replace(" ", "")

    for key, css in _FONT_MAP.items():
        if key.replace(" ", "") in lower:
            return css

    # Fall back: quote the original name and add a generic family
    base = re.sub(r"(PS)?MT$", "", clean)
    base = re.sub(r"-(Bold|Italic|Regular|Light|Medium|Semibold|Heavy).*", "", base)
    if "serif" in lower and "sans" not in lower:
        return f'"{base.strip()}", serif'
    return f'"{base.strip()}", sans-serif'


def _int_color_to_hex(color_int: int) -> str:
    r = (color_int >> 16) & 0xFF
    g = (color_int >> 8) & 0xFF
    b = color_int & 0xFF
    return f"#{r:02x}{g:02x}{b:02x}"


# ---------------------------------------------------------------------------
# Default styles (used as fallback)
# ---------------------------------------------------------------------------
def default_styles() -> dict:
    return {
        "name_font": '"Helvetica Neue", Helvetica, Arial, sans-serif',
        "name_size": "22pt",
        "name_color": "#1a1a2e",
        "name_align": "center",
        "contact_font": '"Helvetica Neue", Helvetica, Arial, sans-serif',
        "contact_size": "9pt",
        "contact_color": "#555555",
        "heading_font": '"Helvetica Neue", Helvetica, Arial, sans-serif',
        "heading_size": "11pt",
        "heading_color": "#1a1a2e",
        "heading_border": True,
        "body_font": '"Helvetica Neue", Helvetica, Arial, sans-serif',
        "body_size": "10pt",
        "body_color": "#222222",
        "accent_color": "#1a1a2e",
        "bullet_color": "#333333",
        "subheading_color": "#555555",
        "margin_top": "0.6in",
        "margin_right": "0.7in",
        "margin_bottom": "0.6in",
        "margin_left": "0.7in",
        "line_height": "1.35",
    }


# ---------------------------------------------------------------------------
# PDF extraction (PyMuPDF / fitz)
# ---------------------------------------------------------------------------
def extract_styles_from_pdf(file_obj) -> dict:
    """Analyze an uploaded PDF and return a style dict."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.warning("PyMuPDF not installed — using default styles")
        return default_styles()

    try:
        file_obj.seek(0)
        doc = fitz.open(stream=file_obj.read(), filetype="pdf")
    except Exception:
        logger.exception("Failed to open PDF for style extraction")
        return default_styles()

    try:
        return _analyze_pdf(doc)
    except Exception:
        logger.exception("Style extraction failed — falling back to defaults")
        return default_styles()
    finally:
        doc.close()


def _analyze_pdf(doc) -> dict:
    """Internal: walk the first page(s) and extract font/color/margin info."""
    styles = default_styles()

    # Collect spans from the first page (and second if available for more data)
    spans = []
    pages_to_scan = min(len(doc), 2)
    page = doc[0]
    page_width = page.rect.width
    page_height = page.rect.height

    for page_idx in range(pages_to_scan):
        p = doc[page_idx]
        blocks = p.get_text("dict", flags=0)["blocks"]
        for block in blocks:
            if "lines" not in block:
                continue
            for line in block["lines"]:
                for span in line["spans"]:
                    text = span["text"].strip()
                    if not text:
                        continue
                    spans.append({
                        "text": text,
                        "font": span["font"],
                        "size": round(span["size"], 1),
                        "color": span["color"],
                        "flags": span["flags"],
                        "bbox": span["bbox"],
                        "page": page_idx,
                    })

    if not spans:
        return styles

    # --- Categorize by font size ---
    size_counter = Counter(s["size"] for s in spans)
    body_size = size_counter.most_common(1)[0][0]  # most common = body

    distinct_sizes = sorted(set(s["size"] for s in spans), reverse=True)
    max_size = distinct_sizes[0]

    # Name = largest text on page 0
    name_spans = [s for s in spans if s["page"] == 0 and abs(s["size"] - max_size) < 0.5]

    # Headings = larger-than-body text (but not name-sized)
    heading_size = body_size
    for sz in distinct_sizes:
        if sz < max_size - 1 and sz > body_size + 0.5:
            heading_size = sz
            break

    heading_spans = [s for s in spans if abs(s["size"] - heading_size) < 1.0 and s["size"] > body_size + 0.5]
    body_spans = [s for s in spans if abs(s["size"] - body_size) < 0.5]

    # --- Fonts ---
    body_font_counts = Counter(s["font"] for s in body_spans)
    body_font = body_font_counts.most_common(1)[0][0] if body_font_counts else "Helvetica"

    name_font = name_spans[0]["font"] if name_spans else body_font
    heading_font = (
        Counter(s["font"] for s in heading_spans).most_common(1)[0][0]
        if heading_spans
        else body_font
    )

    styles["body_font"] = _map_font(body_font)
    styles["body_size"] = f"{body_size:.0f}pt"
    styles["name_font"] = _map_font(name_font)
    styles["name_size"] = f"{max_size:.0f}pt"
    styles["heading_font"] = _map_font(heading_font)
    styles["heading_size"] = f"{heading_size:.0f}pt"
    styles["contact_font"] = _map_font(body_font)
    styles["contact_size"] = f"{max(8, body_size - 1):.0f}pt"

    # --- Colors ---
    body_color_counts = Counter(s["color"] for s in body_spans)
    if body_color_counts:
        styles["body_color"] = _int_color_to_hex(body_color_counts.most_common(1)[0][0])

    if name_spans:
        styles["name_color"] = _int_color_to_hex(name_spans[0]["color"])

    heading_color_counts = Counter(s["color"] for s in heading_spans)
    if heading_color_counts:
        styles["heading_color"] = _int_color_to_hex(heading_color_counts.most_common(1)[0][0])

    # Accent color: look for non-black, non-body colors (often used in headings or lines)
    body_hex = styles["body_color"]
    accent_candidates = Counter()
    for s in spans:
        c = _int_color_to_hex(s["color"])
        if c not in ("#000000", body_hex, "#ffffff", "#fefefe"):
            accent_candidates[c] += 1
    if accent_candidates:
        styles["accent_color"] = accent_candidates.most_common(1)[0][0]
    else:
        styles["accent_color"] = styles["heading_color"]

    styles["bullet_color"] = styles["body_color"]
    styles["subheading_color"] = styles["body_color"]

    # --- Check for heading underlines (drawing objects) ---
    drawings = page.get_drawings()
    has_heading_lines = False
    if heading_spans and drawings:
        for d in drawings:
            if d["type"] == "l":  # line
                # Check if it's a horizontal line near a heading
                y = d["rect"][1]
                for hs in heading_spans:
                    if abs(y - hs["bbox"][3]) < 8:  # within 8pt below heading
                        has_heading_lines = True
                        break
            if has_heading_lines:
                break
    styles["heading_border"] = has_heading_lines

    # --- Margins (from text bounding boxes on first page) ---
    page0_spans = [s for s in spans if s["page"] == 0]
    if page0_spans:
        min_x = min(s["bbox"][0] for s in page0_spans)
        max_x = max(s["bbox"][2] for s in page0_spans)
        min_y = min(s["bbox"][1] for s in page0_spans)

        styles["margin_left"] = f"{max(0.3, min_x / 72):.2f}in"
        styles["margin_right"] = f"{max(0.3, (page_width - max_x) / 72):.2f}in"
        styles["margin_top"] = f"{max(0.3, min_y / 72):.2f}in"
        styles["margin_bottom"] = f"{max(0.3, 0.5):.2f}in"

    # --- Name alignment ---
    if name_spans:
        name_left = min(s["bbox"][0] for s in name_spans)
        name_right = max(s["bbox"][2] for s in name_spans)
        name_center = (name_left + name_right) / 2
        page_center = page_width / 2
        if abs(name_center - page_center) < page_width * 0.1:
            styles["name_align"] = "center"
        else:
            styles["name_align"] = "left"

    # --- Line height ---
    # Estimate from body text line spacing
    if len(body_spans) >= 2:
        ys = sorted(set(s["bbox"][1] for s in body_spans if s["page"] == 0))
        if len(ys) >= 2:
            gaps = [ys[i + 1] - ys[i] for i in range(min(10, len(ys) - 1))]
            avg_gap = sum(gaps) / len(gaps)
            lh = avg_gap / body_size if body_size > 0 else 1.35
            styles["line_height"] = f"{max(1.1, min(1.8, lh)):.2f}"

    return styles


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------
def extract_styles_from_docx(file_obj) -> dict:
    """Analyze an uploaded DOCX and return a style dict."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return default_styles()

    styles = default_styles()

    try:
        file_obj.seek(0)
        doc = Document(io.BytesIO(file_obj.read()))
    except Exception:
        logger.exception("Failed to open DOCX for style extraction")
        return styles

    # Collect font info from runs
    font_info = []  # list of (text, font_name, size_pt, color_hex, bold)
    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text.strip()
            if not text:
                continue
            font = run.font
            name = font.name
            size = font.size.pt if font.size else None
            color = None
            if font.color and font.color.rgb:
                color = f"#{font.color.rgb}"
            bold = font.bold
            font_info.append({
                "text": text,
                "font": name,
                "size": size,
                "color": color,
                "bold": bold,
            })

    if not font_info:
        return styles

    # Filter entries with known sizes
    sized = [f for f in font_info if f["size"] is not None]
    if not sized:
        # Try to at least get font name
        named = [f for f in font_info if f["font"]]
        if named:
            common_font = Counter(f["font"] for f in named).most_common(1)[0][0]
            css_font = _map_font(common_font)
            styles["body_font"] = css_font
            styles["name_font"] = css_font
            styles["heading_font"] = css_font
            styles["contact_font"] = css_font
        return styles

    # Body size = most common
    size_counter = Counter(round(f["size"], 1) for f in sized)
    body_size = size_counter.most_common(1)[0][0]

    distinct_sizes = sorted(set(round(f["size"], 1) for f in sized), reverse=True)
    max_size = distinct_sizes[0]

    # Heading size
    heading_size = body_size
    for sz in distinct_sizes:
        if sz < max_size - 0.5 and sz > body_size + 0.5:
            heading_size = sz
            break

    # Fonts
    body_entries = [f for f in sized if abs(f["size"] - body_size) < 0.5 and f["font"]]
    if body_entries:
        common_body = Counter(f["font"] for f in body_entries).most_common(1)[0][0]
        styles["body_font"] = _map_font(common_body)
        styles["contact_font"] = styles["body_font"]

    name_entries = [f for f in sized if abs(f["size"] - max_size) < 0.5 and f["font"]]
    if name_entries:
        styles["name_font"] = _map_font(name_entries[0]["font"])

    heading_entries = [f for f in sized if abs(f["size"] - heading_size) < 1.0 and f["size"] > body_size + 0.5 and f["font"]]
    if heading_entries:
        styles["heading_font"] = _map_font(Counter(f["font"] for f in heading_entries).most_common(1)[0][0])

    styles["body_size"] = f"{body_size:.0f}pt"
    styles["name_size"] = f"{max_size:.0f}pt"
    styles["heading_size"] = f"{heading_size:.0f}pt"
    styles["contact_size"] = f"{max(8, body_size - 1):.0f}pt"

    # Colors
    colored = [f for f in font_info if f["color"]]
    body_colored = [f for f in colored if f["size"] and abs(f["size"] - body_size) < 0.5]
    if body_colored:
        styles["body_color"] = Counter(f["color"] for f in body_colored).most_common(1)[0][0]

    name_colored = [f for f in colored if f["size"] and abs(f["size"] - max_size) < 0.5]
    if name_colored:
        styles["name_color"] = name_colored[0]["color"]

    heading_colored = [f for f in colored if f["size"] and abs(f["size"] - heading_size) < 1.0 and f["size"] > body_size + 0.5]
    if heading_colored:
        styles["heading_color"] = Counter(f["color"] for f in heading_colored).most_common(1)[0][0]

    # Accent: non-black colors
    body_hex = styles["body_color"]
    accent_cands = Counter()
    for f in colored:
        c = f["color"].lower()
        if c not in ("#000000", body_hex.lower(), "#ffffff"):
            accent_cands[c] += 1
    if accent_cands:
        styles["accent_color"] = accent_cands.most_common(1)[0][0]

    # Check for heading borders in paragraph borders
    styles["heading_border"] = False
    for para in doc.paragraphs:
        if para.runs and any(r.font.size and r.font.size.pt and r.font.size.pt > body_size + 0.5 for r in para.runs if r.font.size):
            pPr = para._p.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
            if pPr is not None:
                pBdr = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr")
                if pBdr is not None:
                    styles["heading_border"] = True
                    break

    # Name alignment
    for para in doc.paragraphs:
        if para.runs and any(r.font.size and r.font.size.pt and abs(r.font.size.pt - max_size) < 0.5 for r in para.runs if r.font.size):
            if para.alignment is not None:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    styles["name_align"] = "center"
                elif para.alignment == WD_ALIGN_PARAGRAPH.LEFT:
                    styles["name_align"] = "left"
                elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    styles["name_align"] = "right"
            break

    return styles
