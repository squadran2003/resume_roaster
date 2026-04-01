"""
Rewrite a DOCX file in-place, replacing text while preserving original formatting.

Strategy:
1. Open the original DOCX.
2. Walk all paragraphs (including those inside table cells).
3. Match paragraphs to the structured rewrite JSON by section/entry order.
4. Replace paragraph text while copying style from the first original run.
5. Unmatched paragraphs are left unchanged (safe default).
"""

import copy
import io
import logging
import re
from difflib import SequenceMatcher

from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# Minimum similarity ratio to consider a paragraph a match
_MATCH_THRESHOLD = 0.35


def rewrite_docx_preserving_format(original_file, rewrite_json: dict) -> bytes:
    """
    Open `original_file` (a Django File / S3 file), apply the structured
    rewrite JSON, and return the modified DOCX as bytes.
    """
    original_file.seek(0)
    doc = Document(io.BytesIO(original_file.read()))

    # Build a flat list of replacement texts from the structured JSON
    replacements = _build_replacement_list(rewrite_json)

    # Gather all paragraphs: body + table cells
    all_paragraphs = _collect_paragraphs(doc)

    # Match and replace
    _apply_replacements(all_paragraphs, replacements)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_replacement_list(data: dict) -> list[dict]:
    """
    Flatten the structured JSON into a list of replacement items.
    Each item has a 'type' (name, contact, summary, section_title, heading,
    subheading, bullet) and 'text' to write.
    """
    items = []

    if data.get("name"):
        items.append({"type": "name", "text": data["name"]})
    if data.get("contact"):
        items.append({"type": "contact", "text": data["contact"].replace("|", "  |  ")})
    if data.get("summary"):
        items.append({"type": "summary", "text": data["summary"]})

    for section in data.get("sections", []):
        items.append({"type": "section_title", "text": section.get("title", "")})
        for entry in section.get("entries", []):
            if entry.get("heading"):
                items.append({"type": "heading", "text": entry["heading"]})
            if entry.get("subheading"):
                items.append({"type": "subheading", "text": entry["subheading"]})
            for bullet in entry.get("bullets", []):
                items.append({"type": "bullet", "text": bullet})

    return items


def _collect_paragraphs(doc: Document) -> list:
    """Collect all paragraphs from the document body and tables."""
    paragraphs = []

    for para in doc.paragraphs:
        paragraphs.append(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    paragraphs.append(para)

    return paragraphs


def _normalize(text: str) -> str:
    """Normalize text for comparison: lowercase, collapse whitespace, strip bullets."""
    text = text.strip().lower()
    text = re.sub(r'^[\u2022\-\*\u25cf\u25cb\u2023\u25aa\u25b8]\s*', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text


def _similarity(a: str, b: str) -> float:
    """Return similarity ratio between two strings."""
    if not a or not b:
        return 0.0
    return SequenceMatcher(None, _normalize(a), _normalize(b)).ratio()


def _is_section_header(para_text: str, replacement_text: str) -> bool:
    """Check if a paragraph looks like a section header matching the replacement."""
    norm_para = _normalize(para_text)
    norm_repl = _normalize(replacement_text)
    # Exact or near-exact match on short uppercase-style text
    if len(norm_para) < 60 and _similarity(norm_para, norm_repl) > 0.6:
        return True
    return False


def _apply_replacements(paragraphs: list, replacements: list[dict]):
    """
    Walk through paragraphs and replacements sequentially.
    For each replacement, find the best matching paragraph (looking ahead)
    and replace its text while preserving formatting.
    """
    para_idx = 0
    total_paras = len(paragraphs)

    for repl in replacements:
        repl_text = repl["text"]
        if not repl_text.strip():
            continue

        # Look ahead from current position to find best match
        best_score = 0.0
        best_idx = -1
        # Search window: look ahead up to 15 paragraphs, or more for early items
        search_limit = min(para_idx + 20, total_paras)

        for i in range(para_idx, search_limit):
            para_text = paragraphs[i].text.strip()
            if not para_text:
                continue

            if repl["type"] == "section_title":
                if _is_section_header(para_text, repl_text):
                    best_idx = i
                    best_score = 1.0
                    break
            else:
                score = _similarity(para_text, repl_text)
                if score > best_score:
                    best_score = score
                    best_idx = i

        if best_idx >= 0 and best_score >= _MATCH_THRESHOLD:
            _replace_paragraph_text(paragraphs[best_idx], repl_text)
            para_idx = best_idx + 1
        else:
            # No match found — skip this replacement (safe default)
            logger.debug(
                "DOCX rewriter: no match for %s '%s' (best_score=%.2f)",
                repl["type"], repl_text[:50], best_score,
            )


def _replace_paragraph_text(paragraph, new_text: str):
    """
    Replace the text content of a paragraph while preserving the formatting
    of the first run.
    """
    if not paragraph.runs:
        # No runs — just set the text directly
        paragraph.text = new_text
        return

    # Save formatting from the first run
    first_run = paragraph.runs[0]
    saved_rpr = copy.deepcopy(first_run._r.get_or_add_rPr())

    # Clear all existing runs
    for run in paragraph.runs:
        run._r.getparent().remove(run._r)

    # Add a single new run with preserved formatting
    from docx.oxml import OxmlElement

    new_r = OxmlElement("w:r")
    new_r.append(saved_rpr)
    new_t = OxmlElement("w:t")
    new_t.set(qn("xml:space"), "preserve")
    new_t.text = new_text
    new_r.append(new_t)
    paragraph._p.append(new_r)
